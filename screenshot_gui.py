import tkinter as tk
from tkinter import filedialog, messagebox
import os
import csv
import datetime
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
from webdriver_manager.chrome import ChromeDriverManager
from pathlib import Path
import time
from docx import Document
from docx.shared import Inches
import re

def safe_filename(url):
    return re.sub(r'[^a-zA-Z0-9_-]', '_', url)[:50]

def take_screenshots_and_generate_docx(csv_path, output_base_folder, template_path):
    try:
        with open(csv_path, newline='') as csvfile:
            reader = csv.reader(csvfile)
            urls = [cell.strip() for row in reader for cell in row if cell.strip()]

        date_folder = datetime.date.today().isoformat()
        output_folder = Path(output_base_folder) / date_folder
        output_folder.mkdir(parents=True, exist_ok=True)

        chrome_options = Options()
        chrome_options.add_argument("--headless")
        chrome_options.add_argument("--window-size=1920,1080")
        service = Service(ChromeDriverManager().install())
        driver = webdriver.Chrome(service=service, options=chrome_options)

        screenshots = []
        for idx, url in enumerate(urls, start=1):
            try:
                driver.get(url)
                time.sleep(2)
                filename = f"{safe_filename(url)}.png"
                screenshot_path = output_folder / filename
                driver.save_screenshot(str(screenshot_path))
                screenshots.append((f"$URL{idx}", f"$IMG{idx}", url, screenshot_path))
            except Exception as e:
                print(f"Error capturing {url}: {e}")

        driver.quit()

        old_doc = Document(template_path)
        new_doc = Document()

        for para in old_doc.paragraphs:
            original_text = para.text
            new_text = original_text
            insert_image_after = None

            for url_tag, img_tag, real_url, image_path in screenshots:
                if url_tag in new_text:
                    new_text = new_text.replace(url_tag, real_url)
                if img_tag in new_text:
                    new_text = new_text.replace(img_tag, "")
                    insert_image_after = image_path

            if new_text.strip() or not insert_image_after:
                new_doc.add_paragraph(new_text)

            if insert_image_after:
                new_doc.add_paragraph().add_run().add_picture(str(insert_image_after), width=Inches(6))

        output_docx = output_folder / f"Report_{date_folder}.docx"
        new_doc.save(output_docx)
        messagebox.showinfo("Done", f"Report saved to: {output_docx}")

    except Exception as e:
        messagebox.showerror("Error", str(e))

def select_file(var, filetypes):
    file_path = filedialog.askopenfilename(filetypes=filetypes)
    if file_path:
        var.set(file_path)

def select_folder(var):
    folder_path = filedialog.askdirectory()
    if folder_path:
        var.set(folder_path)

def run_tool():
    csv_path = csv_var.get()
    output_folder = folder_var.get()
    template_path = template_var.get()
    if not csv_path or not output_folder or not template_path:
        messagebox.showwarning("Missing Input", "Please select the CSV file, output folder, and DOCX template.")
        return
    take_screenshots_and_generate_docx(csv_path, output_folder, template_path)

root = tk.Tk()
root.title("Magic Screenshot Report Tool")

csv_var = tk.StringVar()
folder_var = tk.StringVar()
template_var = tk.StringVar()

tk.Label(root, text="Select URLs.csv:").grid(row=0, column=0, sticky="e")
tk.Entry(root, textvariable=csv_var, width=50).grid(row=0, column=1)
tk.Button(root, text="Browse", command=lambda: select_file(csv_var, [("CSV files", "*.csv")])).grid(row=0, column=2)

tk.Label(root, text="Select Output Folder:").grid(row=1, column=0, sticky="e")
tk.Entry(root, textvariable=folder_var, width=50).grid(row=1, column=1)
tk.Button(root, text="Browse", command=lambda: select_folder(folder_var)).grid(row=1, column=2)

tk.Label(root, text="Select DOCX Template:").grid(row=2, column=0, sticky="e")
tk.Entry(root, textvariable=template_var, width=50).grid(row=2, column=1)
tk.Button(root, text="Browse", command=lambda: select_file(template_var, [("Word Documents", "*.docx")])).grid(row=2, column=2)

tk.Button(root, text="Generate Report", command=run_tool, bg="green", fg="white").grid(row=3, column=1, pady=10)

tk.Label(root, text="Created by Florin Baicu - May 2025", fg="gray").grid(row=4, column=1, pady=5)

root.mainloop()
